import requests
import nmap
import re
import ipinfo
from docx import Document
import json

def get_public_ip():
    try:
        response = requests.get("https://api64.ipify.org?format=json")
        if response.status_code == 200:
            return response.json()["ip"]
        else:
            return None
    except Exception as e:
        print(f"Error getting public IP address: {e}")
        return None

def get_ipinfo_details(api_key, ip_address):
    handler = ipinfo.getHandler(api_key)
    details = handler.getDetails(ip_address)
    return details.all

def sync_nmap_scan(ip_address, ports="-F"):
    cve_list = []

    try:
        nmScan = nmap.PortScanner()
        nmScan.scan(ip_address, ports)

        for host in nmScan.all_hosts():
            print(f"Open ports for {host}:")
            for proto in nmScan[host].all_protocols():
                lport = nmScan[host][proto].keys()
                for port in lport:
                    service_name = nmScan[host][proto][port]["name"]
                    print(f"Port {port}/{proto.upper()} - Service: {service_name}")

                    # Perform a script scan for vulnerabilities
                    res = nmScan.scan(hosts=ip_address, arguments=f"--script=vuln -p {port}")

                    # Extract and print CVEs
                    cve_list.extend(re.findall(r'(CVE-\d+-\d+)', str(res)))

                    if cve_list:
                        print(f"CVEs found for Port {port}/{proto.upper()}: {', '.join(cve_list)}")
                    else:
                        print(f"No CVEs found for Port {port}/{proto.upper()}")

        return cve_list

    except nmap.PortScannerError as e:
        print(f"Nmap PortScannerError: {e}")
        return []
    except Exception as e:
        print(f"An unexpected error occurred: {e}")
        return []

def search_otx(cve_list):
    otx_api_key = '7b6dd71042e43c475d56af1424af56331dc2b2e8ee6ab6426471843b2ce72bd6'
    otx_results = {}
    for cve in cve_list:
        otx_api_url = f'https://otx.alienvault.com/api/v1/indicators/cve/{cve}/'
        headers = {'X-OTX-API-KEY': otx_api_key}

        try:
            response = requests.get(otx_api_url, headers=headers)
            if response.status_code == 200:
                data = response.json()
                description = data.get('description', 'N/A')
                mitre_url = data.get('mitre_url', 'N/A')

                otx_results[cve] = {
                    'description': description,
                    'mitre_url': mitre_url,
                    
                    # Add more fields as needed
                }
                """ print(f"OTX Information for CVE {cve}:")
                print("----------------------------------------------------------------")
                print("Description:", description)
                print("----------------------------------------------------------------")
                print("Mitre URL:", mitre_url)
                print("----------------------------------------------------------------")
 """
                print("\n")
               
                
            else:
                print(f"Failed to retrieve OTX information for CVE {cve}. Status code: {response.status_code}")
        except requests.RequestException as e:
            print(f"Error during OTX request: {e}")

    otx_json = json.dumps(otx_results, indent=2)
    
    print(otx_json)
    return otx_json

def generate_report(public_ip, ipinfo_details, port_results, otx_results):
    document = Document()
    document.add_heading('vulnerability_scan_report', level=1)

    # Public IP and IPinfo details
    document.add_heading('Public IP Information', level=2)
    document.add_paragraph(f"Public IP: {public_ip}")
    document.add_paragraph(f"IPinfo Details: {ipinfo_details}")
    
    # Nmap Scan Results
    document.add_heading('Nmap Scan Results', level=2)
    for result in port_results:
        document.add_paragraph(result)

    # OTX Results
    document.add_heading('OTX Scan Results', level=2)
    otx_results = json.loads(otx_results)

    for cve, details in otx_results.items():
        document.add_heading(f' {cve}', level=3)
        document.add_paragraph(f'Description: {details["description"]}')
        document.add_paragraph(f'Mitre URL: {details["mitre_url"]}')
       

    # Save the document
    document.save('vulnerability_scan_report.docx')
    print("Report generated successfully.")


def main():
    
    ip_address = input("Enter an IP address: ")
    public_ip = get_public_ip()
    if public_ip:
        print(f"Your public IP address is: {public_ip}")

        # Get information for the public IP address
        access_token = 'cb6b9ce8ce2be2'
        ipinfo_details = get_ipinfo_details(access_token, public_ip)

        # Specify the range of ports or use "-F" for common ports
        ports_to_scan = input("Enter ports to scan (e.g., 1-1000, -F for common ports): ")

        # Perform the Nmap port scan and retrieve CVEs
        cve_results = sync_nmap_scan(ip_address, ports_to_scan)

        # Search for additional information on OTX
        otx_results = search_otx(cve_results)

        # Generate and save the report
        generate_report(public_ip, ipinfo_details, cve_results, otx_results)
    else:
        print("Failed to retrieve public IP address.")

if __name__ == "__main__":
    main()
